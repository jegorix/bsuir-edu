#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

try:
    import tkinter as tk
    from tkinter import filedialog
except Exception:  # pragma: no cover - tkinter can be unavailable in headless setups
    tk = None
    filedialog = None

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

SERIF_FONT = "Noto Serif"
CODE_FONT = "JetBrains Mono"
SERIF_SIZE = 12
CODE_SIZE = 8
PROJECT_ROOT = Path(__file__).resolve().parent


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Создает DOCX-листинг для всех .c файлов в выбранной папке "
            "и ее подпапках."
        )
    )
    parser.add_argument(
        "source_dir",
        nargs="?",
        help="Путь к папке с .c файлами. Если не указан, откроется выбор папки.",
    )
    parser.add_argument(
        "-o",
        "--output",
        help=(
            "Путь к выходному .docx файлу. По умолчанию файл создается "
            "в выбранной папке."
        ),
    )
    return parser.parse_args()


def choose_directory(initial_dir: Path) -> Path | None:
    if tk is None or filedialog is None:
        return None

    try:
        root = tk.Tk()
        root.withdraw()
        root.update()
        selected = filedialog.askdirectory(
            initialdir=str(initial_dir),
            title="Выберите папку с .c файлами",
        )
        root.destroy()
    except Exception:
        return None

    if not selected:
        return None

    return Path(selected).expanduser().resolve()


def resolve_source_dir(source_dir_arg: str | None) -> Path:
    if source_dir_arg:
        source_dir = Path(source_dir_arg).expanduser().resolve()
    else:
        source_dir = choose_directory(Path.cwd())
        if source_dir is None:
            entered = input("Введите путь к папке с .c файлами: ").strip()
            if not entered:
                raise SystemExit("Папка не выбрана.")
            source_dir = Path(entered).expanduser().resolve()

    if not source_dir.exists():
        raise SystemExit(f"Папка не найдена: {source_dir}")
    if not source_dir.is_dir():
        raise SystemExit(f"Это не папка: {source_dir}")

    return source_dir


def resolve_output_path(source_dir: Path, output_arg: str | None) -> Path:
    if output_arg:
        output_path = Path(output_arg).expanduser().resolve()
        if output_path.suffix.lower() != ".docx":
            output_path = output_path.with_suffix(".docx")
    else:
        output_dir = build_default_output_dir(source_dir)
        output_path = output_dir / f"{source_dir.name}_listing.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    return output_path


def build_default_output_dir(source_dir: Path) -> Path:
    match = re.search(r"lab[-_ ]?(\d+)", source_dir.name, flags=re.IGNORECASE)
    if match:
        return PROJECT_ROOT / "LISTING" / f"LAB{match.group(1)}"

    return PROJECT_ROOT / "LISTING" / source_dir.name


def collect_c_files(source_dir: Path) -> list[Path]:
    return sorted(
        (path for path in source_dir.rglob("*.c") if path.is_file()),
        key=lambda path: str(path.relative_to(source_dir)),
    )


def build_display_names(source_dir: Path, c_files: list[Path]) -> dict[Path, str]:
    name_counts: dict[str, int] = {}
    for file_path in c_files:
        name_counts[file_path.name] = name_counts.get(file_path.name, 0) + 1

    display_names: dict[Path, str] = {}
    for file_path in c_files:
        if name_counts[file_path.name] == 1:
            display_names[file_path] = file_path.name
        else:
            display_names[file_path] = str(file_path.relative_to(source_dir))

    return display_names


def read_source_text(file_path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1251"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue

    return file_path.read_text(encoding="utf-8", errors="replace")


def set_run_font(run, font_name: str, size: int, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold

    run_properties = run._element.get_or_add_rPr()
    font_element = run_properties.find(qn("w:rFonts"))
    if font_element is None:
        font_element = OxmlElement("w:rFonts")
        run_properties.append(font_element)

    for font_key in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        font_element.set(qn(font_key), font_name)


def add_text_paragraph(
    document: Document,
    text: str,
    *,
    bold: bool = False,
    centered: bool = False,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, SERIF_FONT, SERIF_SIZE, bold=bold)


def add_code_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.right_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    run = paragraph.add_run(text.expandtabs(4))
    set_run_font(run, CODE_FONT, CODE_SIZE)


def build_document(source_dir: Path, c_files: list[Path], output_path: Path) -> None:
    document = Document()
    display_names = build_display_names(source_dir, c_files)

    add_text_paragraph(document, "ПРИЛОЖЕНИЕ А", bold=True, centered=True)
    add_text_paragraph(document, "(листинг кода)", centered=True)
    document.add_paragraph()

    for index, file_path in enumerate(c_files):
        add_text_paragraph(document, display_names[file_path])

        source_lines = read_source_text(file_path).splitlines()
        if not source_lines:
            add_code_paragraph(document, "")
        else:
            for line in source_lines:
                add_code_paragraph(document, line)

        if index != len(c_files) - 1:
            document.add_paragraph()

    document.save(output_path)


def main() -> int:
    args = parse_args()
    source_dir = resolve_source_dir(args.source_dir)
    output_path = resolve_output_path(source_dir, args.output)
    c_files = collect_c_files(source_dir)

    if not c_files:
        print(f"В папке {source_dir} не найдено ни одного .c файла.", file=sys.stderr)
        return 1

    build_document(source_dir, c_files, output_path)

    print(f"Найдено .c файлов: {len(c_files)}")
    print(f"Документ создан: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
